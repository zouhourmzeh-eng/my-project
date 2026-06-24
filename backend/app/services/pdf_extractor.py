import os
import re
import logging
import fitz  # PyMuPDF
import docx  # python-docx
import openpyxl  # openpyxl
import olefile
from striprtf.striprtf import rtf_to_text

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts raw text from a PDF file using PyMuPDF.
    """
    if not file_path or not os.path.exists(file_path):
        logger.warning(f"File not found: {file_path}")
        return ""
    
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
    
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts raw text from a DOCX file using python-docx.
    """
    if not file_path or not os.path.exists(file_path):
        return ""
    
    text = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                text.append(" | ".join(row_cells))
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
    
    return "\n".join(text).strip()

def extract_text_from_xlsx(file_path: str) -> str:
    """
    Extracts raw text from a XLSX file using openpyxl.
    """
    if not file_path or not os.path.exists(file_path):
        return ""
    
    text = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(str(cell) for cell in row if cell is not None)
                if row_str.strip():
                    text.append(row_str)
    except Exception as e:
        logger.error(f"Error extracting text from XLSX {file_path}: {e}")
    
    return "\n".join(text).strip()

def extract_text_from_doc(file_path: str) -> str:
    """
    Extracts text from a legacy .doc file.
    Supports RTF (Rich Text Format) and OLE2 binary word streams.
    """
    if not file_path or not os.path.exists(file_path):
        return ""

    # 1. Check if the file is RTF (Rich Text Format) disguised as .doc
    try:
        with open(file_path, "rb") as f:
            header = f.read(10)
            if header.startswith(b"{\\rtf"):
                f.seek(0)
                try:
                    content_str = f.read().decode("utf-8", errors="ignore")
                except Exception:
                    f.seek(0)
                    content_str = f.read().decode("cp1252", errors="ignore")
                return rtf_to_text(content_str).strip()
    except Exception as e:
        logger.warning(f"Error checking RTF header for {file_path}: {e}")

    # 2. Check if the file is OLE binary
    if not olefile.isOleFile(file_path):
        logger.warning(f"File is neither RTF nor a valid OLE file: {file_path}")
        return ""

    try:
        with olefile.OleFileIO(file_path) as ole:
            if not ole.exists('WordDocument'):
                logger.warning(f"No WordDocument stream found in OLE file: {file_path}")
                return ""
            
            with ole.openstream('WordDocument') as stream:
                data = stream.read()

            # Decode characters. In Word .doc format, text is stored as CP1252 or UTF-16LE.
            # We extract continuous runs of printable text to bypass binary structures.
            
            # UTF-16LE decoding runs (trying offset 0 and offset 1 for alignment safety)
            text_utf16_0 = data.decode('utf-16-le', errors='ignore')
            printable_utf16_0 = re.findall(
                r'[A-Za-z0-9À-ÿ\s\-\.\,\?\!\:\;\(\)\[\]\{\}\'\"\/\\\|\=\+\*\&\%\$\#\@\_]{4,}', 
                text_utf16_0
            )
            
            text_utf16_1 = data[1:].decode('utf-16-le', errors='ignore') if len(data) > 1 else ""
            printable_utf16_1 = re.findall(
                r'[A-Za-z0-9À-ÿ\s\-\.\,\?\!\:\;\(\)\[\]\{\}\'\"\/\\\|\=\+\*\&\%\$\#\@\_]{4,}', 
                text_utf16_1
            )

            # Choose the best UTF-16LE run
            len_utf16_0 = sum(len(s) for s in printable_utf16_0)
            len_utf16_1 = sum(len(s) for s in printable_utf16_1)
            if len_utf16_0 >= len_utf16_1:
                printable_utf16 = printable_utf16_0
                utf16_len = len_utf16_0
            else:
                printable_utf16 = printable_utf16_1
                utf16_len = len_utf16_1
            
            # CP1252 (ANSI) decoding runs
            text_cp1252 = data.decode('cp1252', errors='ignore')
            printable_cp1252 = re.findall(
                r'[A-Za-z0-9À-ÿ\s\-\.\,\?\!\:\;\(\)\[\]\{\}\'\"\/\\\|\=\+\*\&\%\$\#\@\_]{8,}', 
                text_cp1252
            )
            cp1252_len = sum(len(s) for s in printable_cp1252)

            if utf16_len > cp1252_len:
                raw_text = "\n".join(printable_utf16)
            else:
                raw_text = "\n".join(printable_cp1252)

            # Strip leading/trailing spaces and remove empty lines
            lines = []
            for line in raw_text.splitlines():
                line = line.strip()
                if line:
                    lines.append(line)
            
            return "\n".join(lines)

    except Exception as e:
        logger.error(f"Error extracting text from OLE .doc {file_path}: {e}")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from PDF, DOCX, DOC, or XLSX file based on file extension.
    """
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".doc":
        return extract_text_from_doc(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_text_from_xlsx(file_path)
    else:
        logger.warning(f"Unsupported file format for text extraction: {ext}")
        return ""
